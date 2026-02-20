"""Fixtures for ingestion tests: generates test files in various formats."""

from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture
def ingest_fixtures(tmp_path: Path) -> Path:
    """Create a temporary directory with test documents in multiple formats."""
    # Plain text files (5 documents to meet minimum threshold)
    for i in range(1, 6):
        (tmp_path / f"doc_{i:02d}.txt").write_text(
            f"Document {i} content.\n\n"
            f"This is the body of document number {i}. It contains several sentences "
            f"that provide enough text for basic analysis. The document covers generic "
            f"topics related to compliance and regulatory frameworks.\n\n"
            f"Section two of document {i} continues with additional material about "
            f"organizational policies and best practices for documentation."
        )

    # Markdown file
    (tmp_path / "article.md").write_text(
        "# Sample Article\n\n"
        "This is a **bold** introduction to the article.\n\n"
        "## Section One\n\n"
        "Content with [a link](https://example.com) and `inline code`.\n\n"
        "## Section Two\n\n"
        "More content with *italic* text and a list:\n\n"
        "- Item one\n"
        "- Item two\n"
        "- Item three\n"
    )

    # HTML file
    (tmp_path / "page.html").write_text(
        "<html><head><title>Test Page</title></head><body>"
        "<nav>Navigation menu</nav>"
        "<article>"
        "<h1>Main Article</h1>"
        "<p>This is the main content of the article. It discusses important topics.</p>"
        "<p>Second paragraph with more details about the subject matter.</p>"
        "</article>"
        "<footer>Copyright notice</footer>"
        "</body></html>"
    )

    return tmp_path


@pytest.fixture
def pdf_fixture(tmp_path: Path) -> Path:
    """Create a minimal PDF fixture using PyMuPDF."""
    import fitz

    pdf_path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This is a test PDF document.\n\nIt has two paragraphs.")
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture
def docx_fixture(tmp_path: Path) -> Path:
    """Create a minimal DOCX fixture using python-docx."""
    from docx import Document as DocxDocument

    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It has two paragraphs of content.")
    doc.save(str(docx_path))
    return docx_path
